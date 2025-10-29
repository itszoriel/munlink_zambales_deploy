"""Utilities for rendering DOCX digital documents from templates.

Renders municipality-specific templates from public/digital_docs_template
with fallbacks and embeds a QR code image for verification.
"""
from __future__ import annotations

import os
from pathlib import Path
from datetime import datetime, timedelta
from typing import Optional, Tuple, Any, Dict
import json

from flask import current_app
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Inches, Pt, RGBColor
from docx import Document as WordDocument
import qrcode


def _slugify(name: str) -> str:
    return name.lower().replace(' ', '-').replace('_', '-').replace('–', '-').replace('—', '-')


def _municipality_alias(name: str) -> str:
    # Handle common aliasing between folder names and canonical names
    aliases = {
        'Santa Cruz': 'Sta Cruz',
        'Sta. Cruz': 'Sta Cruz',
        'SanAntonio': 'San Antonio',
        'SanAntonio ': 'San Antonio',
    }
    return aliases.get(name, name)


def _resolve_template_path(base_dir: Path, municipality_name: str, document_code: str) -> Optional[Path]:
    alias_name = _municipality_alias(municipality_name)
    slug = _slugify(alias_name)

    # 1) municipality/{code}.docx
    # 2) municipality-slug/{code}.docx
    # 3) municipality/{municipality}.docx
    # 4) municipality-slug/{slug}.docx
    # 5) any .docx in municipality folder (pick first)
    # 6) _default/{code}.docx
    # 7) _default/{municipality}.docx
    # 8) any .docx in _default (pick first)
    candidates = [
        base_dir / alias_name / f"{document_code}.docx",
        base_dir / slug / f"{document_code}.docx",
        base_dir / alias_name / f"{alias_name}.docx",
        base_dir / slug / f"{slug}.docx",
    ]
    for p in candidates:
        if p.exists():
            return p

    for folder in [base_dir / alias_name, base_dir / slug]:
        if folder.exists() and folder.is_dir():
            any_docs = sorted(folder.glob("*.docx"))
            for p in any_docs:
                if p.is_file():
                    return p

    default_candidates = [
        base_dir / "_default" / f"{document_code}.docx",
        base_dir / "_default" / f"{alias_name}.docx",
        base_dir / "_default" / f"{slug}.docx",
    ]
    for p in default_candidates:
        if p.exists():
            return p

    default_folder = base_dir / "_default"
    if default_folder.exists() and default_folder.is_dir():
        any_default_docs = sorted(default_folder.glob("*.docx"))
        for p in any_default_docs:
            if p.is_file():
                return p

    return None


def _ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def _build_qr_image(temp_dir: Path, qr_data: str) -> Path:
    _ensure_dir(temp_dir)
    img = qrcode.make(qr_data)
    out = temp_dir / "qr.png"
    img.save(out)
    return out


def _load_municipal_meta(base_dir: Path, municipality_name: str) -> Dict[str, Any]:
    alias_name = _municipality_alias(municipality_name)
    slug = _slugify(alias_name)
    for folder in [base_dir / alias_name, base_dir / slug, base_dir / "_default"]:
        meta_path = folder / "meta.json"
        if meta_path.exists():
            try:
                return json.loads(meta_path.read_text(encoding="utf-8"))
            except Exception:
                return {}
    return {}


def render_request_docx(*, request, document_type, user) -> Tuple[Path, Path]:
    """
    Render a DOCX for a document request using municipality-specific template.

    Returns: (docx_path, pdf_target_path) where pdf_target_path is the desired
    output path for the converted PDF (not yet created).
    """
    # Base dirs
    project_root = Path(current_app.config.get('BASE_DIR') or Path(__file__).resolve().parents[3])
    templates_dir = project_root / 'public' / 'digital_docs_template'

    municipality_name = getattr(getattr(request, 'municipality', None), 'name', '') or str(request.municipality_id)
    municipality_slug = _slugify(municipality_name)

    template_path = _resolve_template_path(templates_dir, municipality_name, document_type.code)
    if not template_path:
        raise FileNotFoundError(f"Template not found for {municipality_name} / {document_type.code}")

    # Output directory per request
    upload_base = Path(current_app.config['UPLOAD_FOLDER'])
    out_dir = upload_base / 'municipalities' / municipality_slug / 'documents' / 'generated' / str(request.id)
    _ensure_dir(out_dir)

    # QR data and image
    qr_base = current_app.config.get('QR_BASE_URL', 'http://localhost:3000/verify')
    qr_payload = f"{qr_base}?req={request.request_number}"
    qr_img_path = _build_qr_image(out_dir, qr_payload)

    # Build context
    resident_full_name = ' '.join(filter(None, [getattr(user, 'first_name', None), getattr(user, 'last_name', None)])) or getattr(user, 'username', 'Resident')
    barangay_name = getattr(getattr(request, 'barangay', None), 'name', '')
    issue_date = datetime.utcnow()
    valid_until = issue_date + timedelta(days=max(int(getattr(document_type, 'processing_days', 3)), 1))

    meta = _load_municipal_meta(templates_dir, municipality_name) or {}

    ctx = {
        'resident_full_name': resident_full_name,
        'resident_name': resident_full_name,            # alias
        'resident_address': request.delivery_address or '',
        'address': request.delivery_address or '',      # alias
        'municipality_name': municipality_name,
        'municipality': municipality_name,              # alias
        'barangay_name': barangay_name,
        'document_name': document_type.name,
        'document': document_type.name,                 # alias
        'document_code': document_type.code,
        'request_number': request.request_number,
        'request_no': request.request_number,           # alias
        'purpose': request.purpose or '',
        'issue_date': issue_date.strftime('%Y-%m-%d'),
        'date': issue_date.strftime('%B %d, %Y'),       # alias (long)
        'date_issued': issue_date.strftime('%B %d, %Y'),
        'valid_until': valid_until.strftime('%Y-%m-%d'),
        'meta': meta,
    }

    # Render
    doc = DocxTemplate(str(template_path))
    # Helpful filters
    try:
        doc.jinja_env.filters['date_long'] = lambda s: datetime.strptime(s, '%Y-%m-%d').strftime('%B %d, %Y') if isinstance(s, str) else (s.strftime('%B %d, %Y') if isinstance(s, datetime) else s)
        doc.jinja_env.filters['currency'] = lambda n: f"₱{float(n):,.2f}" if n is not None else ''
        doc.jinja_env.filters['upper'] = lambda x: (x or '').upper()
    except Exception:
        pass
    # Insert QR image as InlineImage if placeholder exists in template
    # We always provide it; template can ignore it if unused
    qr_inline = InlineImage(doc, str(qr_img_path), width=Mm(30))
    ctx['qr_image'] = qr_inline
    ctx['qr'] = qr_inline  # alias
    # Optional seal/signature from meta
    try:
        seal_file = meta.get('seal') if isinstance(meta, dict) else None
        logos_base = Path(current_app.config.get('BASE_DIR') or Path(__file__).resolve().parents[3]) / 'public' / 'logos'
        search_folders = [
            templates_dir / _municipality_alias(municipality_name),
            templates_dir / _slugify(municipality_name),
            templates_dir / "_default",
            logos_base / 'municipalities' / _municipality_alias(municipality_name),
            logos_base / 'municipalities' / _slugify(municipality_name),
            logos_base / 'zambales',
        ]
        if seal_file:
            for folder in search_folders:
                fp = folder / seal_file
                if fp.exists():
                    ctx['seal_image'] = InlineImage(doc, str(fp), width=Mm(24))
                    break
        if isinstance(meta, dict) and isinstance(meta.get('signatories'), list) and meta['signatories']:
            sig_file = meta['signatories'][0].get('signature')
            if sig_file:
                for folder in search_folders:
                    fp = folder / sig_file
                    if fp.exists():
                        ctx['signature_image'] = InlineImage(doc, str(fp), width=Mm(30))
                        break
    except Exception:
        pass
    doc.render(ctx)

    docx_out = out_dir / 'document.docx'
    doc.save(str(docx_out))

    # Safety net: append default content if template didn't render key fields
    try:
        wd = WordDocument(str(docx_out))
        full_text = []
        for p in wd.paragraphs:
            full_text.append(p.text or '')
        doc_text = "\n".join(full_text)
        needs_body = True
        # If request number or resident name appears, assume body exists
        if isinstance(request.request_number, str) and request.request_number in doc_text:
            needs_body = False
        if resident_full_name and resident_full_name in doc_text:
            needs_body = False
        if document_type.name and document_type.name in doc_text:
            needs_body = False
        if needs_body:
            # Start on a clean page to avoid first-page overlays from letterhead
            try:
                wd.add_page_break()
            except Exception:
                pass
            title = wd.add_paragraph()
            run = title.add_run(document_type.name)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Calibri'
            title.alignment = 1
            p_intro = wd.add_paragraph(f"This document is issued by {municipality_name}.")
            for r in p_intro.runs:
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0, 0, 0)
                r.font.name = 'Calibri'
            table = wd.add_table(rows=5, cols=2)
            table.style = 'Table Grid'
            rows = [
                ('Request Number', request.request_number),
                ('Issued To', resident_full_name),
                ('Purpose', request.purpose or ''),
                ('Issue Date', issue_date.strftime('%B %d, %Y')),
                ('Valid Until', valid_until.strftime('%B %d, %Y')),
            ]
            for i, (k, v) in enumerate(rows):
                table.cell(i, 0).text = k
                table.cell(i, 1).text = str(v)
                for para in table.cell(i, 0).paragraphs + table.cell(i, 1).paragraphs:
                    for r in para.runs:
                        r.font.size = Pt(11)
                        r.font.color.rgb = RGBColor(0, 0, 0)
                        r.font.name = 'Calibri'
            if isinstance(meta, dict) and isinstance(meta.get('signatories'), list) and meta['signatories']:
                sig = meta['signatories'][0]
                wd.add_paragraph('')
                p_sig_name = wd.add_paragraph(sig.get('name', 'Authorized Official'))
                p_sig_title = wd.add_paragraph(sig.get('title', ''))
                for para in [p_sig_name, p_sig_title]:
                    for r in para.runs:
                        r.font.size = Pt(11)
                        r.font.color.rgb = RGBColor(0, 0, 0)
                        r.font.name = 'Calibri'
            try:
                wd.add_paragraph('Scan to verify:')
                wd.add_picture(str(qr_img_path), width=Inches(1.2))
            except Exception:
                pass
            wd.save(str(docx_out))
    except Exception:
        pass

    pdf_target = out_dir / 'document.pdf'
    return docx_out, pdf_target


